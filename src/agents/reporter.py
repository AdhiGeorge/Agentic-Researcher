from swarm import Agent
from . import get_prompt
import os
import logging

try:
    from fpdf import FPDF
except ImportError:
    FPDF = None
try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)

def build_structured_report(context_variables):
    """
    Build a single, structured report with Introduction, Theory, Formula, Python Code, Usage Example, and Sources.
    """
    intro = context_variables.get('intro', '')
    theory = context_variables.get('theory', '')
    formula = context_variables.get('formula', '')
    usage = context_variables.get('usage', '')
    sources = context_variables.get('sources', '')
    code = context_variables.get('generated_code', '')
    # Fallbacks: try to extract from formatted_research or answer if not present
    if not intro:
        intro = context_variables.get('answer', '')
    if not theory:
        theory = context_variables.get('formatted_research', '')
    if not formula:
        formula = context_variables.get('formula', '')
    if not usage:
        usage = context_variables.get('usage', '')
    if not sources:
        sources = context_variables.get('sources', '')
    # Build the report
    report = f"""
# Volatility Index (VIX): Explanation, Formula, and Python Code

## Introduction
{intro}

## Theory/Background
{theory}

## Mathematical Formula
{formula}

## Python Code
```python
{code}
```

## Usage Example
{usage}

## Sources
{sources}
"""
    return report

def generate_report(context_variables):
    """
    Generate a comprehensive report summarizing the project and export to PDF, TXT, DOCX, or code file as requested.
    Preserves Markdown structure and code blocks as best as possible.
    """
    # Build a single, structured report
    report_content = build_structured_report(context_variables)
    code_content = context_variables.get('generated_code', '')
    export_format = context_variables.get('export_format', 'txt').lower()
    export_dir = context_variables.get('export_path', 'exports')
    os.makedirs(export_dir, exist_ok=True)
    base_name = context_variables.get('export_name', 'research_report')
    exported_files = []
    try:
        # Export main report
        if export_format == 'pdf':
            if not FPDF:
                logger.error('FPDF is not installed. Cannot export PDF.')
                context_variables['project_report'] = 'Export failed: FPDF (pdf export library) is not installed.'
                return context_variables['project_report']
            if not report_content.strip():
                logger.error('No report content to export to PDF.')
                context_variables['project_report'] = 'Export failed: No report content to export.'
                return context_variables['project_report']
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font('Arial', '', 12)
            lines = report_content.split('\n')
            in_code = False
            for line in lines:
                if line.strip().startswith('```'):
                    in_code = not in_code
                    if in_code:
                        pdf.set_font('Courier', '', 10)
                    else:
                        pdf.set_font('Arial', '', 12)
                    continue
                if in_code:
                    pdf.multi_cell(0, 7, line)
                else:
                    pdf.multi_cell(0, 10, line)
            pdf_path = os.path.abspath(os.path.join(export_dir, f"{base_name}.pdf"))
            pdf.output(pdf_path)
            exported_files.append(pdf_path)
        elif export_format == 'docx' and Document:
            doc = Document()
            doc.add_heading('Research Report', 0)
            lines = report_content.split('\n')
            in_code = False
            code_block = []
            for line in lines:
                if line.strip().startswith('```'):
                    in_code = not in_code
                    if not in_code and code_block:
                        doc.add_paragraph('\n'.join(code_block), style='Code')
                        code_block = []
                    continue
                if in_code:
                    code_block.append(line)
                else:
                    doc.add_paragraph(line)
            docx_path = os.path.abspath(os.path.join(export_dir, f"{base_name}.docx"))
            doc.save(docx_path)
            exported_files.append(docx_path)
        elif export_format == 'txt':
            txt_path = os.path.abspath(os.path.join(export_dir, f"{base_name}.txt"))
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(report_content)
            exported_files.append(txt_path)
        # Export code if present
        if code_content:
            code_ext = context_variables.get('code_ext', 'py')
            code_path = os.path.abspath(os.path.join(export_dir, f"{base_name}_code.{code_ext}"))
            with open(code_path, 'w', encoding='utf-8') as f:
                f.write(code_content)
            exported_files.append(code_path)
        if not exported_files:
            context_variables['project_report'] = 'Export failed: No files were generated.'
            return context_variables['project_report']
        context_variables['project_report'] = f"Exported files: {exported_files}"
        return context_variables['project_report']
    except Exception as e:
        logger.error(f"Error exporting report: {e}")
        context_variables['project_report'] = f"Export failed: {e}"
        return context_variables['project_report']

reporter_agent = Agent(
    name="Reporter Agent",
    instructions=get_prompt('reporter'),
    functions=[generate_report],
) 